"""
Document text extractor for PDF and DOCX files.
"""
from __future__ import annotations
import io


def extract_text_from_file(uploaded_file) -> tuple[str, int]:
    """
    Extract plain text from a Streamlit UploadedFile (PDF or DOCX).
    Returns (text, page_count).
    """
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if name.endswith(".pdf"):
        return _from_pdf(raw)
    elif name.endswith(".docx"):
        return _from_docx(raw)
    else:
        raise ValueError(f"Unsupported format: {uploaded_file.name}")


def _from_pdf(raw: bytes) -> tuple[str, int]:
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                pages_text.append(txt)
        return "\n\n".join(pages_text), len(pages_text)
    except ImportError:
        # Fallback: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            texts = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(texts), len(texts)
        except ImportError:
            raise ImportError("Install pdfplumber or pypdf: pip install pdfplumber")


def _from_docx(raw: bytes) -> tuple[str, int]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n\n".join(paragraphs)
        return text, len(paragraphs)
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")
